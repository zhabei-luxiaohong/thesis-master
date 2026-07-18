#!/usr/bin/env python3
from docx import Document
from docx.shared import Pt, Cm
from docx.oxml.ns import qn
import re

INPUT = '/Users/op/WorkBuddy/论文优化/mba_thesis_full.txt'
OUTPUT = '/Users/op/WorkBuddy/论文优化/MBA论文_优化版.docx'

with open(INPUT, 'r', encoding='utf-8') as f:
    content = f.read()

doc = Document()
s = doc.styles['Normal']
s.font.name = '宋体'; s.font.size = Pt(12)
s._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
s.paragraph_format.line_spacing = 1.5

lines = content.split('\n')
h1_set = set()
h2_set = set()
h3_set = set()

# Collect headings
for line in lines:
    t = line.strip()
    if t.startswith('第') and '章' in t and len(t) < 30: h1_set.add(t)
    elif re.match(r'^\d+\.\d+\s', t) and len(t) < 60: 
        if t.count('.') == 1: h2_set.add(t)
        elif t.count('.') >= 2: h3_set.add(t)

in_table = False
table_lines = []

for line in lines:
    t = line.strip()
    if not t:
        if in_table and table_lines:
            # Build table
            rows = [l.split('|') for l in table_lines]
            max_cols = max(len(r) for r in rows)
            tbl = doc.add_table(rows=len(rows), cols=max_cols)
            tbl.style = 'Table Grid'
            for ri, row in enumerate(rows):
                for ci, val in enumerate(row):
                    if ci < max_cols:
                        c = tbl.cell(ri, ci); c.text = ''
                        r = c.paragraphs[0].add_run(val.strip())
                        r.font.size = Pt(9); r.font.name = '宋体'
                        r._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
                        if ri == 0: r.bold = True
            doc.add_paragraph()
            table_lines = []
        in_table = False
        continue
    
    # Table detection
    if '|' in t and not t.startswith('数据来源') and t.count('|') >= 2:
        in_table = True
        table_lines.append(t)
        continue
    
    # Heading
    if t in h1_set:
        doc.add_heading(t, level=1)
    elif t in h2_set:
        doc.add_heading(t, level=2)
    elif t in h3_set:
        doc.add_heading(t, level=3)
    else:
        p = doc.add_paragraph()
        p.paragraph_format.first_line_indent = Cm(0.74)
        p.paragraph_format.line_spacing = 1.5
        r = p.add_run(t)
        r.font.size = Pt(12); r.font.name = '宋体'
        r._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

doc.save(OUTPUT)
print(f'Done: {OUTPUT}')
